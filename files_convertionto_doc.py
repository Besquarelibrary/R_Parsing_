EXTRACTION OF HEADINGS FROM THE DOC FILE


from docx import Document
import re

input_docx_file = "C://Users//DELL//Downloads//Anusha_Updated (1).docx"
section_headings = []

doc = Document(input_docx_file)
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text:
        if re.search(r'^\s*[\d\s]*[A-Z][A-Z\s\-:]*$', text):
            section_headings.append(text)

print("Extracted Section Headings:")
for heading in section_headings:
    print(heading)



--------------------------------------------------------------------------
EXTRACTING THE HEADINGS FROM DOC AND ADD IT IN THE JSON FILE

import json
from docx import Document
import re

input_docx_file = "C://Users//DELL//Downloads//Anusha_Updated (1).docx"
section_headings = []
current_heading = None
content = ""

doc = Document(input_docx_file)
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text:
        if re.search(r'^\s*[\d\s]*[A-Z][A-Z\s\-:]*$', text):
            if current_heading:
                section_headings.append((current_heading, content))
            current_heading = text
            content = ""
        else:
            content += text + "\n"

# Add the last section
if current_heading:
    section_headings.append((current_heading, content))

# Create a dictionary from the section_headings list
sections_dict = {heading: content for heading, content in section_headings}

# Write the dictionary to a JSON file
output_json_file = "C://Users//DELL//Downloads//Anusha.json"
with open(output_json_file, "w") as json_file:
    json.dump(sections_dict, json_file, indent=4)

print(f"JSON file '{output_json_file}' created successfully.")
------------------------------------------------------------------------------

CONVERT THE JSON TO DOC BY EXTRACTING THE HEADINGS


import json
from docx import Document
from docx.shared import Pt  # Import the Pt object
import re

# Read the JSON file
input_json_file = "C://Users//DELL//Downloads//Anusha.json"
with open(input_json_file, "r") as json_file:
    sections_data = json.load(json_file)

# Create a new Word document
output_docx_file = "C://Users//DELL//Downloads//Anusha.docx"
output_doc = Document()

# Iterate through the sections data and add headings and content to the document
for heading, content in sections_data.items():
    # Add heading
    heading_para = output_doc.add_paragraph()
    heading_run = heading_para.add_run(heading)
    heading_run.bold = True
    heading_run.font.size = Pt(14)  # Using Pt from docx.shared

    # Add content
    content_para = output_doc.add_paragraph(content)

    # Add space between sections
    output_doc.add_paragraph()

# Save the new Word document
output_doc.save(output_docx_file)

print(f"Word document '{output_docx_file}' created successfully.")

----------------------------------------------------------------------------------



CREATE THE JSON FROM THE PDF BY EXTRACTING THE HEADINGS

import fitz  # PyMuPDF library
import re
import json

input_pdf_file = "C://Users//DELL//Downloads//Anusha_Updated (1).pdf"
sections_data = {}
current_heading = None
content = ""

pdf_document = fitz.open(input_pdf_file)
for page_num in range(pdf_document.page_count):
    page = pdf_document[page_num]
    text = page.get_text().strip()

    # Split text into lines and process each line
    lines = text.split("\n")
    for i in range(len(lines)):
        line = lines[i].strip()
        if line:
            if re.search(r'^\s*[\d\s]*[A-Z][A-Z\s\-:]*$', line):
                if current_heading:
                    sections_data[current_heading] = content
                current_heading = line
                content = ""
            else:
                content += line + "\n"

# Add the last section
if current_heading:
    sections_data[current_heading] = content

pdf_document.close()

# Save the sections data as a JSON file
output_json_file = "C://Users//DELL//Downloads//anusha(2)_pdf.json"
with open(output_json_file, "w") as json_file:
    json.dump(sections_data, json_file, indent=4)

print(f"JSON file '{output_json_file}' created successfully.")
-----------------------------------------------------------------------------
JSON TO DOC WITH HEADINGS EXTRACTION

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

# Read the JSON file
input_json_file = "C://Users//DELL//Downloads//anusha(2)_pdf.json"
with open(input_json_file, "r") as json_file:
    sections_data = json.load(json_file)

# Create a new Word document
output_docx_file = "C://Users//DELL//Downloads//modified(5)_anusha_pd.docx"
output_doc = Document()

# Iterate through the sections data and add headings and content to the document
for heading, content in sections_data.items():
    # Add heading
    heading_para = output_doc.add_paragraph(heading)
    heading_run = heading_para.runs[0]
    heading_run.bold = True
    heading_run.font.size = Pt(14)  # Adjust font size if needed
    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment to center

    # Add content
    content_para = output_doc.add_paragraph(content)
    content_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set alignment to left

    # Add space between sections
    output_doc.add_paragraph()

# Save the new Word document
output_doc.save(output_docx_file)

print(f"Word document '{output_docx_file}' created successfully.")
-----------------------------------------------------------------------------------
CODE FOR TEMPLATE CREATION USING JINJA2

from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('Resume Template', level=1)

# Add placeholders for skills, project info, education, experience, email, and phone number
placeholders = ['skills', 'project info', 'education', 'experience', 'email', 'phonenumber']
for placeholder in placeholders:
    doc.add_heading(placeholder.capitalize(), level=2)
    doc.add_paragraph(f"{{{{ {placeholder} }}}}")

# Save the template to a Word file
template_docx_file = "C://Users//DELL//Downloads//resume_template.docx"
doc.save(template_docx_file)

print(f"Template created and saved to '{template_docx_file}'.")
