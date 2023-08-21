import os
import json
import fitz  # PyMuPDF library for PDF extraction
from docx import Document
from docx.shared import Pt  # Import the Pt object
import re
import shutil
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# PDF extraction function
def extract_pdf_content(input_pdf_file):
    # ... (same as before)
    section_headings = []

    pdf_document = fitz.open(input_pdf_file)
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text = page.get_text().strip()

        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if line:
                if re.search(r'^\s*[\d\s]*[A-Z][A-Z\s\-:]*$', line):
                    section_headings.append((line, ""))  # Add empty content for now

                # Append content to the last heading
                elif section_headings:
                    section_headings[-1] = (section_headings[-1][0], section_headings[-1][1] + line + "\n")

    pdf_document.close()

    return section_headings


# Docx extraction function
def extract_docx_content(input_docx_file):
    section_headings = []
    current_heading = None
    content = ""

    doc = Document(input_docx_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            if re.search(r'^\s*[\d\s]*[A-Z][A-Z\s\-:]*$', text):
                if current_heading:
                    section_headings.append((current_heading, content))
                current_heading = text
                content = ""
            else:
                content += text + "\n"

    # Add the last section
    if current_heading:
        section_headings.append((current_heading, content))

    return section_headings
    

# Create JSON file from section headings
def create_json_file(section_headings, output_json_file):
    # ... (same as before)
    sections_dict = {heading: content for heading, content in section_headings}

    with open(output_json_file, "w") as json_file:
        json.dump(sections_dict, json_file, indent=4)

    print(f"JSON file '{output_json_file}' created successfully.")


# Process files in a folder
def process_files_in_folder(folder_path):
    # Create a new folder to store JSON files
    new_folder_path = os.path.join(folder_path, "new_json_files")
    os.makedirs(new_folder_path, exist_ok=True)

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            section_headings = extract_pdf_content(file_path)
        elif filename.endswith(".docx"):
            section_headings = extract_docx_content(file_path)
        else:
            continue  # Skip non-PDF and non-docx files

        # Create JSON file for the document
        output_json_filename = f"{os.path.splitext(filename)[0]}.json"
        output_json_file = os.path.join(new_folder_path, output_json_filename)
        create_json_file(section_headings, output_json_file)

def convert_json_to_docx(input_json_file, output_docx_folder):
    with open(input_json_file, "r") as json_file:
        sections_dict = json.load(json_file)

    for heading, content in sections_dict.items():
        output_docx_file = os.path.join(output_docx_folder, f"{os.path.splitext(os.path.basename(input_json_file))[0]}.docx")
        output_doc = Document()

        for heading, content in sections_dict.items():
            heading_para = output_doc.add_paragraph(heading)
            heading_run = heading_para.runs[0]
            heading_run.bold = True
            heading_run.font.size = Pt(14)
            heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            content_para = output_doc.add_paragraph(content)
            content_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            output_doc.add_paragraph()

        output_doc.save(output_docx_file)
        print(f"Word document '{output_docx_file}' created successfully.")


def main():
    input_folder = input("Enter the folder path containing PDF and DOCX files: ")
    new_json_folder = os.path.join(input_folder, "new_json_files")
    new_docx_folder = os.path.join(input_folder, "new_docx_files")

    process_files_in_folder(input_folder)
    
    # Convert JSON files to DOCX files and store in new folder
    os.makedirs(new_docx_folder, exist_ok=True)
    for json_filename in os.listdir(new_json_folder):
        if json_filename.endswith(".json"):
            json_file_path = os.path.join(new_json_folder, json_filename)
            convert_json_to_docx(json_file_path, new_docx_folder)

if __name__ == "__main__":
    main()
