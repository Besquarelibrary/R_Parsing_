import os
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def load_json_file(json_file):
    with open(json_file, "r") as f:
        data = json.load(f)
    return data

def create_template(file_path, json_data):
    doc = Document()

    # Title
    title = doc.add_paragraph("My Resume", style="Title")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    placeholders = {
        "Education": ["EDUCATION DETAILS:", "EDUCATION DETAILS", "EDUCATION:", "EDUCATION", "ACADEMIC DETAILS:", "ACADEMIC DETAILS"],
        "Skills": ["TECHNICAL SKILLS:", "TECHNICAL SKILLS", "SKILLS", "SKILLS:"],
        "Projects": ["PROJECT SUMMARY", "PROJECT SUMMARY:", "PROJECT DETAILS", "PROJECT DETAILS:", "PROJECT INFO", "PROJECT INFO:"],
        "Experience": ["PROFESSIONAL EXPERIENCE:", "PROFESSIONAL EXPERIENCE", "EXPERIENCE", "EXPERIENCE:"]
    }

    for section, placeholder_keys in placeholders.items():
        for key in placeholder_keys:
            if key in json_data:
                placeholder_text = json_data[key]
                doc.add_heading(section, level=1)
                doc.add_paragraph(placeholder_text, style="Normal")

    doc.save(file_path)
    print(f"Template created and saved to {file_path}")

def process_json_files(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.endswith(".json"):
            json_file_path = os.path.join(input_folder, filename)
            json_data = load_json_file(json_file_path)

            template_file_name = filename.replace(".json", "_template.docx")
            template_file_path = os.path.join(output_folder, template_file_name)

            create_template(template_file_path, json_data)

if __name__ == "__main__":
    input_folder_path = "C://Users//DELL//Downloads//test_resume//new_json_files"
    output_folder_path = "C://Users//DELL//Downloads//test_resume//new_json_templates"

    process_json_files(input_folder_path, output_folder_path)
