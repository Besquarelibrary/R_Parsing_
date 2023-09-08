#THIS CODE IS FOR EXTRACTING THE TITLES WHICH ARE IN PASCAL CASE WITH SHADOW AND WITHOUT SHADOW, CAPITAL TEXT WITHOUT SHADOW ALONG WITH FONT SIZE SHOULD BE MORE THAN 12 THEN IT WILL CREATE THE NEW DOCX WITH ADDING THE UPPERCASE AND BOLD TO EXTRACTED TITLES
----------------------------------------------------------------------
import os
import docx
import re

# Specify the folder path containing the resumes
folder_path = 'C:/Users/DELL/Downloads/test_resume/new'

# Define a regular expression pattern to match Title Case headings
title_case_pattern = r'\b[A-Z][a-z]*\b'

# Create a new folder to store the output files
output_folder = 'C:/Users/DELL/Downloads/test_resume/new_mod'
os.makedirs(output_folder, exist_ok=True)

# Iterate through all files in the folder
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        # Construct the full path of the DOCX file
        docx_file_path = os.path.join(folder_path, filename)

        # Initialize variables to keep track of the current heading and its font size
        current_heading = ""
        current_font_size = None

        # Initialize a variable to store the content for the current heading
        current_content = []

        # Open the .docx file using python-docx
        doc = docx.Document(docx_file_path)

        # Create a new DOCX document for the current resume
        output_doc = docx.Document()

        # Iterate through paragraphs in the document
        for paragraph in doc.paragraphs:
            # Check if the paragraph text matches the Title Case pattern
            matches = re.findall(title_case_pattern, paragraph.text)

            # Check if the paragraph has a larger font size
            has_larger_font = False
            for run in paragraph.runs:
                if run.font.size and run.font.size.pt > 12:
                    has_larger_font = True
                    break

            # If matches are found and the font size is larger than 12, assume it's a new heading
            if matches and has_larger_font:
                # Store the previous heading (if it's in Title Case and has a larger font size)
                if current_heading and re.match(title_case_pattern, current_heading) and current_font_size > 12:
                    # Add the title (in uppercase and bold) to the output document
                    title_run = output_doc.add_paragraph(current_heading.upper()).runs[0]
                    title_run.bold = True

                    # Add the content for the previous heading to the output document
                    for content_line in current_content:
                        output_doc.add_paragraph(content_line)

                # Update the current heading, font size, and reset content
                current_heading = ' '.join(matches)
                current_font_size = max(run.font.size.pt for run in paragraph.runs if run.font.size)
                current_content = []
            else:
                # Add the paragraph text to the content for the current heading
                current_content.append(paragraph.text)

        # Store the last heading (if it's in Title Case and has a larger font size)
        if current_heading and re.match(title_case_pattern, current_heading) and current_font_size > 12:
            # Add the title (in uppercase and bold) to the output document
            title_run = output_doc.add_paragraph(current_heading.upper()).runs[0]
            title_run.bold = True

            # Add the content for the last heading to the output document
            for content_line in current_content:
                output_doc.add_paragraph(content_line)

        # Specify the path to save the output DOCX file for the current resume
        output_file_path = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}_Output.docx")

        # Save the output DOCX file for the current resume
        output_doc.save(output_file_path)

        # Print a message indicating the location of the saved file
        print(f"Output for {filename} saved to: {output_file_path}")

# Print a message indicating the location of the output folder
print(f"Output files saved in folder: {output_folder}")